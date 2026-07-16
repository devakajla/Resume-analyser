import fitz  # PyMuPDF
from docx import Document
import easyocr
import os

try:
    import pdfplumber
    pdfplumber_available = True
except ImportError:
    pdfplumber_available = False


reader = easyocr.Reader(['en'], gpu=False)


def detect_format(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    supported = {'.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg'}
    if ext not in supported:
        raise ValueError(f"Unsupported format: {ext}")
    return ext


def extract_from_pdf(file_path):
    # 1. Try to use pdfplumber for layout & table extraction first
    if pdfplumber_available:
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # extract_text maintains horizontal alignment of tables/columns
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If pdfplumber successfully extracted text and it's not a scanned PDF
            if len(text.strip()) >= 50:
                return text
        except Exception:
            pass  # Fallback to PyMuPDF if pdfplumber fails

    # 2. Fallback: PyMuPDF for standard digital PDFs
    try:
        doc = fitz.open(file_path)
    except Exception:
        return ""

    text = ""
    for page in doc:
        page_text = page.get_text()
        text += page_text

    # 3. Fallback: OCR for scanned PDFs (no text layer)
    if len(text.strip()) < 50:
        text = ""
        for page in doc:
            # Convert page to image
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")

            # Save temp image for OCR
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                tmp.write(img_bytes)
                tmp_path = tmp.name

            # OCR the image
            page_text = extract_from_image(tmp_path)
            text += page_text + "\n"

            # Cleanup temp file
            import os as _os
            _os.remove(tmp_path)

    doc.close()
    return text


def extract_from_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join([cell.text for cell in row.cells])
            text += "\n" + row_text

    return text


def extract_from_image(file_path):
    results = reader.readtext(file_path)
    text = "\n".join([result[1] for result in results])
    return text


def extract_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

import re


def clean_text(text):
    """Clean extracted text for better processing."""
    # Remove extra whitespace and newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)

    # Remove common resume artifacts
    text = re.sub(r'Page \d+ of \d+', '', text)
    text = re.sub(r'\x0c', '', text)  # form feed characters

    # Normalize unicode
    text = text.encode('ascii', 'ignore').decode('ascii')

    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)

    # Remove empty lines at start/end
    text = text.strip()

    return text

def parse_resume(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = detect_format(file_path)

    if ext == '.pdf':
        raw_text = extract_from_pdf(file_path)
    elif ext == '.docx':
        raw_text = extract_from_docx(file_path)
    elif ext in ('.png', '.jpg', '.jpeg'):
        raw_text = extract_from_image(file_path)
    elif ext == '.txt':
        raw_text = extract_from_txt(file_path)

    return clean_text(raw_text)



# Quick test
if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        text = parse_resume(sys.argv[1])
        print(f"Extracted {len(text)} characters")
        print(text[:500])



